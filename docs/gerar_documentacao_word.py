from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUTPUT = DOCS / "Pet_Guardian_IoT_Documentacao_Completa.docx"
ASSETS.mkdir(exist_ok=True)

NAVY = "163A5F"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
GOLD = "D99A2B"
RED = "A33A3A"
GREEN = "2B7A4B"
TEXT = "1F2933"
MUTED = "5C6773"

GITHUB = "https://github.com/LauroAnd34/pet-monitor-iot"
LINKS = {
    "Repositório completo": GITHUB,
    "Baixar documentação completa em Word": f"{GITHUB}/raw/refs/heads/master/docs/Pet_Guardian_IoT_Documentacao_Completa.docx",
    "README principal": f"{GITHUB}/blob/master/README.md",
    "Firmware do ESP32 Hub": f"{GITHUB}/blob/master/src/esp32_pet_hub_dual/esp32_pet_hub_dual.ino",
    "Firmware do ESP32 com OV7670": f"{GITHUB}/blob/master/src/esp32_ov7670_non_fifo_node/esp32_ov7670_non_fifo_node.ino",
    "Aplicativo Android": f"{GITHUB}/tree/master/android_pet_guardian_app",
    "Backend Supabase": f"{GITHUB}/tree/master/cloud/supabase",
    "Arquitetura da nuvem": f"{GITHUB}/blob/master/cloud/README_CLOUD_ARCHITECTURE.md",
    "Diagrama Wokwi": f"{GITHUB}/blob/master/wokwi/diagram.json",
    "Esquema elétrico no GitHub": f"{GITHUB}/blob/master/docs/ESQUEMA_ELETRICO_COMPLETO.md",
    "ESP32 - documentação oficial": "https://docs.espressif.com/projects/esp-idf/en/stable/esp32/get-started/",
    "ESP32 - datasheet oficial": "https://www.espressif.com/sites/default/files/documentation/esp32_datasheet_en.pdf",
    "Arduino IDE": "https://docs.arduino.cc/software/ide/",
    "Supabase": "https://supabase.com/docs",
    "Android Studio": "https://developer.android.com/studio",
    "Wokwi": "https://docs.wokwi.com/",
    "DHT11/DHT22 - guia técnico": "https://learn.adafruit.com/dht",
    "HC-SR04 - referência técnica": "https://learn.sparkfun.com/tutorials/hc-sr04-ultrasonic-distance-sensor-hookup-guide",
}


FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def draw_box(draw, xy, text, color=BLUE, fontsize=24):
    draw.rounded_rectangle(xy, radius=20, fill="white", outline=f"#{color}", width=4)
    left, top, right, bottom = xy
    lines = text.split("\n")
    line_height = fontsize + 5
    total_height = len(lines) * line_height
    y = top + ((bottom - top) - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font(fontsize))
        x = left + ((right - left) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=f"#{TEXT}", font=font(fontsize))
        y += line_height


def arrow(draw, start, end, color=BLUE):
    draw.line([start, end], fill=f"#{color}", width=5)
    x2, y2 = end
    x1, y1 = start
    length = max(((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5, 1)
    ux, uy = (x2 - x1) / length, (y2 - y1) / length
    px, py = -uy, ux
    size = 18
    points = [
        (x2, y2),
        (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55),
        (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55),
    ]
    draw.polygon(points, fill=f"#{color}")


def save_architecture():
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 45), "Arquitetura geral do Pet Guardian IoT", anchor="mm", font=font(42, True), fill=f"#{NAVY}")
    draw_box(draw, (60, 245, 460, 475), "Sensores e atuadores\nDHT11, HC-SR04, PIR\ngás, LDR, motor e bomba", GREEN)
    draw_box(draw, (570, 265, 890, 455), "ESP32 Hub\ndos sensores", BLUE, 28)
    draw_box(draw, (1280, 245, 1740, 475), "Supabase\ntelemetria, comandos\ne Storage privado", GOLD, 27)
    draw_box(draw, (570, 700, 890, 890), "ESP32 dedicada\nà OV7670", BLUE, 28)
    draw_box(draw, (60, 680, 460, 910), "Câmera OV7670\nsem FIFO", GREEN, 28)
    draw_box(draw, (1280, 680, 1740, 910), "Aplicativo Android\npainel, controles\nfotos e ajustes", RED, 27)
    arrow(draw, (460, 360), (570, 360), GREEN)
    arrow(draw, (890, 360), (1280, 360), GOLD)
    arrow(draw, (460, 795), (570, 795), GREEN)
    arrow(draw, (890, 795), (1280, 430), GOLD)
    arrow(draw, (1425, 475), (1425, 680), RED)
    arrow(draw, (1595, 680), (1595, 475), BLUE)
    draw.text((1080, 300), "HTTPS / JSON", anchor="mm", font=font(22), fill=f"#{MUTED}")
    draw.text((1080, 650), "BMP via HTTPS", anchor="mm", font=font(22), fill=f"#{MUTED}")
    path = ASSETS / "arquitetura_pet_guardian.png"
    image.save(path)
    return path


def save_hub_wiring():
    image = Image.new("RGB", (1800, 1320), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 45), "Esquema funcional de ligação - ESP32 Hub", anchor="mm", font=font(40, True), fill=f"#{NAVY}")
    draw_box(draw, (700, 530, 1100, 780), "ESP32 HUB\nsensores e atuadores", NAVY, 30)
    left = [
        ((40, 150, 440, 330), "DHT11\nDATA GPIO 4"),
        ((40, 380, 440, 560), "HC-SR04 ração\nTRIG 18 / ECHO 19"),
        ((40, 610, 440, 790), "HC-SR04 água\nTRIG 5 / ECHO 17"),
        ((40, 840, 440, 1020), "PIR\nOUT GPIO 27"),
        ((40, 1070, 440, 1250), "Gás AO 34\nLDR AO 35"),
    ]
    right = [
        ((1360, 150, 1760, 330), "Buzzer\nGPIO 32"),
        ((1360, 380, 1760, 560), "Iluminação\nsinal GPIO 13"),
        ((1360, 610, 1760, 790), "Ponte H alimentador\nIN1 14 / IN2 12 / EN 25"),
        ((1360, 840, 1760, 1020), "Ponte H bomba\nIN3 26 / IN4 23 / EN 33"),
        ((1360, 1070, 1760, 1250), "Fonte externa\n+ GND comum"),
    ]
    for xy, text in left:
        draw_box(draw, xy, text, GREEN, 24)
        arrow(draw, (xy[2], (xy[1] + xy[3]) // 2), (700, 655), GREEN)
    for xy, text in right:
        color = RED if "Ponte" in text else GOLD
        draw_box(draw, xy, text, color, 22)
        arrow(draw, (1100, 655), (xy[0], (xy[1] + xy[3]) // 2), color)
    path = ASSETS / "esquema_hub.png"
    image.save(path)
    return path


def save_camera_wiring():
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 45), "Esquema funcional de ligação - ESP32 com OV7670", anchor="mm", font=font(38, True), fill=f"#{NAVY}")
    draw_box(draw, (40, 410, 480, 740), "ESP32 dedicada\nà câmera", NAVY, 32)
    draw_box(draw, (1320, 410, 1760, 740), "OV7670 sem FIFO\nalimentação: 3,3 V", GREEN, 30)
    signals = [
        ("SIOD 21 / SIOC 22", 170),
        ("VSYNC 34 / HREF 35", 305),
        ("XCLK 32 / PCLK 33", 440),
        ("D0 27 / D1 5 / D2 2 / D3 15", 575),
        ("D4 14 / D5 13 / D6 12 / D7 4", 710),
        ("RESET → 3,3 V / PWDN → GND", 845),
        ("VCC → 3,3 V / GND → GND", 980),
    ]
    for text, y in signals:
        bbox = draw.textbbox((0, 0), text, font=font(23))
        width = bbox[2] - bbox[0] + 55
        draw.rounded_rectangle((900 - width / 2, y - 42, 900 + width / 2, y + 42), radius=18, fill=f"#{LIGHT_GRAY}", outline=f"#{BLUE}", width=3)
        draw.text((900, y), text, anchor="mm", font=font(23), fill=f"#{TEXT}")
        arrow(draw, (480, 575), (900 - width / 2, y), BLUE)
        arrow(draw, (900 + width / 2, y), (1320, 575), GREEN)
    path = ASSETS / "esquema_camera.png"
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run(run, size=11, bold=False, color=TEXT, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_last_image_alt(doc, description):
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set("descr", description)
    inline.docPr.set("title", description)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "PET GUARDIAN IOT  |  DOCUMENTAÇÃO TÉCNICA"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.runs[0], size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Lauro Andrade  |  ")
    set_run(run, size=8.5, color=MUTED)
    add_field(footer, "PAGE")


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), size=27, bold=True, color=NAVY)
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(18)
        set_run(sub.add_run(subtitle), size=14, color=BLUE)


def add_para(doc, text="", bold_prefix=None, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]), italic=italic)
    else:
        set_run(p.add_run(text), italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item))


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(item))


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[idx])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(text)), size=font_size, bold=True, color="FFFFFF")
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[idx])
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    set_run(p.add_run(f"{title}: "), bold=True, color=accent)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_link(doc, label, url, description=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(f"{label}: "), bold=True, color=NAVY)
    add_hyperlink(p, url, url)
    if description:
        set_run(p.add_run(f" — {description}"), color=MUTED)


def build_document():
    arch = save_architecture()
    hub = save_hub_wiring()
    camera = save_camera_wiring()

    doc = Document()
    configure_styles(doc)

    # Capa
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("RELATÓRIO TÉCNICO E GUIA DE REPRODUÇÃO"), size=11, bold=True, color=GOLD)
    add_title(doc, "Pet Guardian IoT", "Sistema de monitoramento e cuidado remoto de animais")
    doc.add_picture(str(arch), width=Inches(6.2))
    set_last_image_alt(doc, "Arquitetura geral do Pet Guardian IoT")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(metadata.add_run("Autor: Lauro Andrade\n"), size=11, bold=True, color=NAVY)
    set_run(metadata.add_run("Internet das Coisas, sistemas embarcados, computação em nuvem e Android\n"), size=10.5, color=MUTED)
    set_run(metadata.add_run("Fortaleza, 2026"), size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(p, "Acessar o repositório completo no GitHub", GITHUB)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_header_footer(section)

    doc.add_heading("Sumário", level=1)
    toc = [
        "1. Resumo executivo",
        "2. Problema e justificativa",
        "3. Objetivos e funcionalidades",
        "4. Arquitetura e funcionamento",
        "5. Componentes e especificações",
        "6. Esquema elétrico completo",
        "7. Software embarcado",
        "8. Backend Supabase e protocolos",
        "9. Aplicativo Android",
        "10. Montagem e instalação passo a passo",
        "11. Operação do sistema",
        "12. Testes, validação e solução de problemas",
        "13. Segurança e limitações",
        "14. Conclusão",
        "15. Links reais e referências",
    ]
    for item in toc:
        add_para(doc, item)
    add_note(doc, "Documento reproduzível", "As pinagens, alturas dos recipientes e caminhos de arquivos apresentados correspondem aos firmwares publicados no repositório do projeto.")

    doc.add_heading("1. Resumo executivo", level=1)
    add_para(doc, "O Pet Guardian IoT é um sistema integrado para acompanhamento e cuidado remoto de animais domésticos. Ele monitora temperatura, umidade, luminosidade, gás, presença, nível de água e nível de ração; controla alimentador, bomba e iluminação; e permite solicitar fotos por um aplicativo Android.")
    add_para(doc, "A solução utiliza duas placas ESP32 independentes. A primeira funciona como hub dos sensores e atuadores. A segunda é dedicada à câmera OV7670 sem FIFO e envia as imagens diretamente ao Supabase. Essa separação evita sobrecarga no hub e torna o fluxo de fotos independente.")
    add_para(doc, "Palavras-chave: ", bold_prefix="Palavras-chave:")

    doc.add_heading("2. Problema e justificativa", level=1)
    add_para(doc, "Quando o tutor está fora de casa, pode não perceber rapidamente que o pet ficou sem água ou ração, que houve alteração ambiental, que o animal se aproximou do equipamento ou que algum dispositivo ficou offline. Sistemas comerciais costumam ser caros, fechados e pouco adaptáveis.")
    add_para(doc, "O projeto resolve esse problema reunindo sensores, comandos, registros e fotografias em uma solução de baixo custo e código aberto. O tutor recebe informações no celular, consulta o histórico e aciona recursos remotamente.")
    add_note(doc, "Decisão de arquitetura", "A câmera não passa pelo ESP32 dos sensores. O nó OV7670 consulta sua própria fila e envia a foto diretamente ao banco/Storage.")

    doc.add_heading("3. Objetivos e funcionalidades", level=1)
    doc.add_heading("3.1 Objetivo geral", level=2)
    add_para(doc, "Desenvolver um sistema IoT capaz de monitorar e auxiliar o cuidado remoto de um animal doméstico por meio de sensores, atuadores, câmera, nuvem e aplicativo Android.")
    doc.add_heading("3.2 Objetivos específicos", level=2)
    add_bullets(doc, [
        "Medir temperatura, umidade, luminosidade, gás e presença.",
        "Calcular o nível de ração de 0 a 100% para um recipiente de 10 cm.",
        "Calcular o nível de água de 0 a 100% para um recipiente de 14 cm.",
        "Controlar motor do alimentador, bomba d'água e iluminação.",
        "Capturar imagens com uma OV7670 sem FIFO.",
        "Armazenar telemetria, comandos, eventos e fotos na nuvem.",
        "Exibir painel, histórico, controles, fotos, notificações e ajustes no Android.",
    ])
    doc.add_heading("3.3 Funcionalidades disponíveis", level=2)
    add_bullets(doc, [
        "Reconexão automática ao Wi-Fi e melodia ao conectar.",
        "Fotos manuais e automáticas por movimento.",
        "Álbuns, favoritos, compartilhamento e salvamento na galeria.",
        "Linha do tempo do pet e comparação diária.",
        "Limpeza automática configurável de fotos antigas.",
        "Notificações Android e status dos dispositivos.",
    ])

    doc.add_heading("4. Arquitetura e funcionamento", level=1)
    doc.add_picture(str(arch), width=Inches(6.3))
    set_last_image_alt(doc, "Fluxo entre sensores, ESP32, Supabase, câmera e aplicativo Android")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figura 1 — Arquitetura geral do Pet Guardian IoT.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_steps(doc, [
        "O ESP32 Hub lê os sensores e calcula os níveis de água e ração.",
        "O hub publica telemetria no Supabase e consulta comandos destinados a ele.",
        "O aplicativo consulta a nuvem e envia comandos de controle ou captura.",
        "A ESP32 da câmera consulta somente a fila de captura de fotos.",
        "A OV7670 gera uma imagem BMP, que é enviada ao Storage privado.",
        "O aplicativo recebe uma URL assinada temporária e baixa a imagem.",
    ])
    add_note(doc, "Protocolos", "Telemetria e comandos usam Wi-Fi, HTTPS e JSON. As fotografias são transferidas como BMP binário por HTTPS.")

    doc.add_heading("5. Componentes e especificações", level=1)
    components = [
        ("2", "ESP32 DevKit", "Hub dos sensores e nó dedicado da câmera", "Lógica de 3,3 V; Wi-Fi integrado"),
        ("1", "OV7670 sem FIFO", "Captura de imagem", "Barramento paralelo RGB565; alimentação 3,3 V"),
        ("1", "DHT11", "Temperatura e umidade", "Saída digital"),
        ("2", "HC-SR04", "Níveis de ração e água", "Alimentação 5 V; ECHO deve ser reduzido para 3,3 V"),
        ("1", "PIR", "Detecção de presença/movimento", "Saída digital"),
        ("1", "Sensor de gás", "Alerta ambiental", "Saída analógica limitada a 3,3 V"),
        ("1", "LDR/módulo", "Luminosidade", "Saída analógica"),
        ("1", "Ponte H", "Aciona motor e bomba", "Usar fonte externa"),
        ("1", "Motor DC", "Liberação de ração", "Não ligar diretamente à ESP32"),
        ("1", "Bomba d'água", "Reposição/controle de água", "Não ligar diretamente à ESP32"),
        ("1", "Buzzer", "Feedback sonoro", "GPIO 32"),
        ("1", "Fonte externa", "Alimentação dos atuadores", "Dimensionar pela corrente de partida"),
    ]
    add_table(doc, ["Qtd.", "Componente", "Função", "Observação técnica"], components, [0.55, 1.45, 2.0, 2.45], 8.8)
    add_note(doc, "Materiais auxiliares", "Protoboard ou placa, jumpers, resistores de 1 kΩ e 2 kΩ para cada divisor de tensão dos HC-SR04, capacitor eletrolítico próximo à ponte H e gabinete.")

    doc.add_heading("6. Esquema elétrico completo", level=1)
    add_note(doc, "Atenção elétrica", "Nunca alimente motor ou bomba diretamente pelos GPIOs. A OV7670 trabalha em 3,3 V. Todas as fontes de lógica devem compartilhar GND.", fill="FDECEC", accent=RED)
    doc.add_heading("6.1 ESP32 Hub de sensores e atuadores", level=2)
    doc.add_picture(str(hub), width=Inches(6.3))
    set_last_image_alt(doc, "Esquema de ligação dos sensores e atuadores ao ESP32 Hub")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figura 2 — Esquema funcional da ESP32 Hub.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    hub_rows = [
        ("DHT11 DATA", "4", "3,3 V ou 5 V", "Pull-up de 10 kΩ se o módulo não possuir"),
        ("HC-SR04 ração TRIG", "18", "5 V", "Recipiente de 10 cm"),
        ("HC-SR04 ração ECHO", "19", "5 V", "Usar divisor para 3,3 V"),
        ("HC-SR04 água TRIG", "5", "5 V", "Recipiente de 14 cm"),
        ("HC-SR04 água ECHO", "17", "5 V", "Usar divisor para 3,3 V"),
        ("PIR OUT", "27", "Conforme módulo", "Saída digital"),
        ("Sensor de gás AO", "34", "Conforme módulo", "Máximo de 3,3 V no GPIO"),
        ("LDR AO", "35", "3,3 V", "Entrada analógica"),
        ("Iluminação / sinal", "13", "Driver externo", "Use transistor ou relé quando necessário"),
        ("Buzzer", "32", "Conforme buzzer", "Use transistor se necessário"),
        ("Alimentador IN1 / IN2 / EN", "14 / 12 / 25", "Ponte H", "Motor com fonte externa"),
        ("Bomba IN3 / IN4 / EN", "26 / 23 / 33", "Ponte H", "Bomba com fonte externa"),
    ]
    add_table(doc, ["Módulo / sinal", "GPIO", "Alimentação", "Observação"], hub_rows, [1.8, 1.0, 1.3, 2.35], 8.5)
    doc.add_heading("6.2 Divisores de tensão dos sensores ultrassônicos", level=2)
    add_para(doc, "O pino ECHO do HC-SR04 normalmente fornece aproximadamente 5 V, enquanto os GPIOs da ESP32 aceitam no máximo 3,3 V. Em cada ECHO, ligue um resistor de 1 kΩ entre o ECHO e o GPIO e um resistor de 2 kΩ entre o GPIO e o GND. Essa combinação reduz a tensão para uma faixa segura.")
    add_para(doc, "Ligação: ECHO do HC-SR04 → resistor de 1 kΩ → GPIO 19 ou 17 → resistor de 2 kΩ → GND.", bold_prefix="Ligação:")

    doc.add_heading("6.3 ESP32 dedicada à câmera OV7670", level=2)
    doc.add_picture(str(camera), width=Inches(6.3))
    set_last_image_alt(doc, "Esquema de ligação entre a ESP32 dedicada e a câmera OV7670 sem FIFO")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figura 3 — Esquema funcional da ESP32 com OV7670.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cam_rows = [
        ("SIOD / SDA", "21"), ("SIOC / SCL", "22"), ("VSYNC", "34"), ("HREF", "35"),
        ("XCLK", "32"), ("PCLK", "33"), ("D0", "27"), ("D1", "5"), ("D2", "2"),
        ("D3", "15"), ("D4", "14"), ("D5", "13"), ("D6", "12"), ("D7", "4"),
        ("RESET", "3,3 V"), ("PWDN", "GND"), ("VCC", "3,3 V"), ("GND", "GND"),
    ]
    add_table(doc, ["Sinal OV7670", "GPIO / ligação ESP32"], cam_rows, [3.1, 3.35], 9.2)
    add_note(doc, "Qualidade da câmera", "Mantenha os fios curtos e organizados. O barramento paralelo da OV7670 é sensível a ruído, mau contato e alimentação instável.")

    doc.add_heading("7. Software embarcado", level=1)
    doc.add_heading("7.1 Firmware do hub", level=2)
    add_para(doc, "O firmware do hub inicializa sensores, atuadores e Wi-Fi. Durante a operação, realiza leituras periódicas, converte as distâncias em percentuais, publica telemetria, consulta comandos e executa ações. Caso a conexão caia, novas tentativas são realizadas. Uma melodia confirma a conexão bem-sucedida.")
    add_link(doc, "Abrir firmware do hub", LINKS["Firmware do ESP32 Hub"])
    doc.add_heading("7.2 Conversão dos níveis", level=2)
    add_para(doc, "O percentual considera a altura útil de cada recipiente. A distância próxima de zero indica recipiente cheio, enquanto a distância próxima da altura total indica recipiente vazio.")
    add_bullets(doc, ["Ração: altura total de 10 cm.", "Água: altura total de 14 cm.", "Resultado exibido no app: valor limitado entre 0 e 100%."])
    doc.add_heading("7.3 Firmware da câmera", level=2)
    add_para(doc, "O firmware dedicado gera XCLK, configura a câmera por SCCB e lê o barramento paralelo. A resolução bruta é 80x60 e a região útil enviada é 72x49. A escolha reduz o consumo de memória e permite manter buffers suficientes para a comunicação HTTPS em uma ESP32 sem PSRAM.")
    add_link(doc, "Abrir firmware da câmera", LINKS["Firmware do ESP32 com OV7670"])

    doc.add_heading("8. Backend Supabase e protocolos", level=1)
    add_para(doc, "O Supabase centraliza telemetria, comandos, metadados de fotos e armazenamento privado das imagens. O hub e a câmera usam filas separadas para impedir que um dispositivo consuma comandos destinados ao outro.")
    protocols = [
        ("Wi-Fi", "Conecta os dispositivos à rede local e à internet.", "Disponível nativamente na ESP32."),
        ("HTTPS", "Transporta telemetria, comandos e imagens.", "Protege tokens e dados durante o transporte."),
        ("JSON", "Representa telemetria e comandos.", "Legível e fácil de integrar."),
        ("BMP binário", "Representa fotografias.", "Simples de produzir sem FIFO e sem PSRAM."),
        ("URL assinada", "Libera acesso temporário à foto privada.", "Evita tornar o bucket público."),
    ]
    add_table(doc, ["Protocolo / formato", "Uso", "Justificativa"], protocols, [1.35, 2.45, 2.65], 9)
    add_steps(doc, [
        "Criar um projeto no Supabase.",
        "Executar cloud/supabase/schema.sql.",
        "Publicar as Edge Functions disponíveis no repositório.",
        "Configurar tokens e segredos fora do código público.",
        "Cadastrar os dispositivos e validar as filas separadas.",
    ])
    add_link(doc, "Documentação da arquitetura da nuvem", LINKS["Arquitetura da nuvem"])
    add_link(doc, "Documentação oficial do Supabase", LINKS["Supabase"])

    doc.add_heading("9. Aplicativo Android", level=1)
    add_para(doc, "O aplicativo Android em Kotlin apresenta painel, histórico, controles, fotos e ajustes. Ele consulta o Supabase, envia comandos, sincroniza fotos e oferece recursos de organização e compartilhamento.")
    add_bullets(doc, [
        "Painel com sensores, níveis e estado do sistema.",
        "Controles do alimentador, bomba e iluminação.",
        "Solicitação remota de fotos.",
        "Álbuns, favoritos, compartilhamento e salvamento na galeria.",
        "Monitoramento inteligente em Ajustes.",
        "Fotos automáticas por movimento, notificações, limpeza, linha do tempo e comparação diária.",
    ])
    add_steps(doc, [
        "Abrir a pasta android_pet_guardian_app no Android Studio.",
        "Configurar AppConfig.kt com os endpoints e tokens adequados.",
        "Compilar com o Android Studio ou executar ./gradlew assembleDebug.",
        "Instalar o APK em um aparelho Android.",
    ])
    add_link(doc, "Abrir código do aplicativo Android", LINKS["Aplicativo Android"])
    add_link(doc, "Baixar/consultar Android Studio", LINKS["Android Studio"])

    doc.add_heading("10. Montagem e instalação passo a passo", level=1)
    doc.add_heading("10.1 Preparação", level=2)
    add_steps(doc, [
        "Separar componentes, conferir tensões e identificar todos os pinos.",
        "Instalar Arduino IDE e o pacote de placas ESP32.",
        "Clonar ou baixar o repositório do GitHub.",
        "Criar o projeto Supabase e preparar as credenciais.",
    ])
    doc.add_heading("10.2 Montagem do hub", level=2)
    add_steps(doc, [
        "Montar apenas ESP32, DHT11, PIR, LDR e gás; validar leituras.",
        "Adicionar os dois HC-SR04 usando divisores nos pinos ECHO.",
        "Conferir a conversão de ração em 10 cm e água em 14 cm.",
        "Adicionar ponte H e fonte externa, mantendo GND comum.",
        "Conectar motor, bomba, iluminação e buzzer; testar um atuador por vez.",
    ])
    doc.add_heading("10.3 Montagem da câmera", level=2)
    add_steps(doc, [
        "Ligar a OV7670 à segunda ESP32 conforme a tabela de pinagem.",
        "Conferir VCC em 3,3 V, RESET em 3,3 V e PWDN em GND.",
        "Usar fios curtos e verificar continuidade.",
        "Gravar o firmware e abrir o Monitor Serial.",
        "Testar /status e /capture.bmp?reason=manual no IP local.",
    ])
    doc.add_heading("10.4 Integração completa", level=2)
    add_steps(doc, [
        "Validar publicação de telemetria no Supabase.",
        "Validar os comandos do hub.",
        "Solicitar uma foto pelo aplicativo.",
        "Confirmar o fluxo app → fila → câmera → Storage → app.",
        "Executar os testes de aceitação da seção seguinte.",
    ])

    doc.add_heading("11. Operação do sistema", level=1)
    add_steps(doc, [
        "Ligar as duas ESP32 e aguardar a conexão Wi-Fi.",
        "Abrir o aplicativo e verificar painel e status.",
        "Usar Controles para acionar alimentador, bomba ou iluminação.",
        "Abrir Fotos e solicitar uma captura.",
        "Usar Ajustes para configurar automações, notificações e limpeza.",
        "Acompanhar histórico, linha do tempo e comparação diária.",
    ])

    doc.add_heading("12. Testes, validação e solução de problemas", level=1)
    tests = [
        ("Inicialização do hub", "Sensores respondem e Wi-Fi reconecta.", "Ver Monitor Serial e telemetria."),
        ("Nível de ração", "0 a 100% com base em 10 cm.", "Medir vazio, metade e cheio."),
        ("Nível de água", "0 a 100% com base em 14 cm.", "Medir vazio, metade e cheio."),
        ("Atuadores", "Somente o dispositivo solicitado é acionado.", "Enviar comandos individualmente."),
        ("Captura local", "Endpoint retorna BMP válido.", "Abrir /capture.bmp?reason=manual."),
        ("Captura remota", "Foto aparece no aplicativo.", "Validar fila, upload e sincronização."),
        ("Aplicativo", "Fotos e Ajustes abrem sem fechar.", "Navegar pelas telas no aparelho."),
    ]
    add_table(doc, ["Teste", "Resultado esperado", "Como verificar"], tests, [1.45, 2.35, 2.65], 8.7)
    doc.add_heading("12.1 Problemas comuns", level=2)
    problems = [
        ("ESP32 não grava", "Reduzir a velocidade de upload para 115200, trocar cabo/porta e manter BOOT pressionado quando necessário."),
        ("Câmera não gera imagem", "Revisar VCC 3,3 V, XCLK, PCLK, VSYNC, HREF, D0-D7 e comprimento dos fios."),
        ("Foto local funciona, mas não aparece no app", "Verificar tokens, URLs, fila exclusiva da câmera, Edge Function e bucket privado."),
        ("Níveis incorretos", "Confirmar alturas de 10 cm e 14 cm e calibrar posição dos HC-SR04."),
        ("ESP32 reinicia ao ligar motor/bomba", "Usar fonte externa adequada, GND comum e capacitor próximo à ponte H."),
    ]
    add_table(doc, ["Problema", "Ação recomendada"], problems, [2.0, 4.45], 9)

    doc.add_heading("13. Segurança e limitações", level=1)
    add_bullets(doc, [
        "Não publicar senhas Wi-Fi, tokens de dispositivos, dashboard_token ou service_role.",
        "Manter o bucket de fotos privado e usar URLs assinadas.",
        "Não aplicar mais de 3,3 V aos GPIOs da ESP32.",
        "Não alimentar motor ou bomba pelos pinos da placa.",
        "Testar a montagem antes de deixá-la funcionando sem supervisão.",
    ])
    add_para(doc, "A OV7670 sem FIFO oferece resolução limitada e exige temporização e cabeamento cuidadosos. Sensores ultrassônicos podem sofrer interferência por formato do recipiente, espuma, inclinação e superfícies irregulares. O protótipo deve ser calibrado e protegido contra água antes do uso contínuo.")

    doc.add_heading("14. Conclusão", level=1)
    add_para(doc, "O Pet Guardian IoT demonstra uma solução completa que une eletrônica, firmware, computação em nuvem e desenvolvimento Android. A arquitetura com duas ESP32 torna o fluxo de câmera independente do hub de sensores e facilita a manutenção. Com este documento e os arquivos públicos do GitHub, outra pessoa pode montar, configurar, testar e evoluir o sistema.")

    doc.add_heading("15. Links reais e referências", level=1)
    add_para(doc, "Os links abaixo são clicáveis e apontam para o código real do projeto ou para documentações técnicas oficiais e guias reconhecidos.")
    for label, url in LINKS.items():
        add_link(doc, label, url)

    doc.core_properties.title = "Pet Guardian IoT - Documentação Completa"
    doc.core_properties.subject = "Relatório técnico, esquema elétrico e guia de reprodução"
    doc.core_properties.author = "Lauro Andrade"
    doc.core_properties.keywords = "IoT, ESP32, OV7670, Android, Supabase, Pet Guardian"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
